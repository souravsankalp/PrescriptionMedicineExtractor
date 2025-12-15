"""
extraction.py

Use PaddleOCR (3.x pipeline) to extract text from an image, group the
words into lines (top-to-bottom, left-to-right), and write them into
a .docx file.

Output path pattern:
    E:\ImageTextSqubix\temp\<file_id>.docx

Requirements in your virtualenv:
    pip install paddlepaddle==3.2.0   # or matching CPU version for your OS
    pip install PaddleOCR
    pip install python-docx
"""

from pathlib import Path

from docx import Document
from paddleocr import PaddleOCR

# ---------------------------------------------------------------------
# Where you want to save .docx files
# ---------------------------------------------------------------------
TEMP_TEXT_DIR = Path(r"E:\ImageTextSqubix\temp")

# ---------------------------------------------------------------------
# Create one global PaddleOCR pipeline object
#   - we disable the heavy doc-orientation and unwarping modules
#     to keep it lighter; you can set lang="en" (Latin + numbers)
# ---------------------------------------------------------------------
ocr = PaddleOCR(
    use_doc_orientation_classify=False,
    use_doc_unwarping=False,
    use_textline_orientation=False,
    lang="en",
)


def _collect_boxes_from_result(result_list):
    """
    Convert PaddleOCR 3.x pipeline output into a list of
    (y_center, x_min, text) tuples for line grouping.

    result_list: returned by ocr.predict(path)
    """
    all_boxes = []

    if not result_list:
        return all_boxes

    # For images we send, there will usually be exactly one page result
    for page in result_list:
        # In 3.x, each item is an OCRResult object with `.res` dict
        res_dict = getattr(page, "res", page)

        # Texts and boxes are stored here
        texts = res_dict.get("rec_texts", [])
        # rec_boxes: [ [x_min, y_min, x_max, y_max], ... ]
        rec_boxes = res_dict.get("rec_boxes", None)
        rec_polys = res_dict.get("rec_polys", None)
        dt_polys = res_dict.get("dt_polys", None)

        for i, text in enumerate(texts):
            text = str(text).strip()
            if not text:
                continue

            box = None

            if rec_boxes is not None and i < len(rec_boxes):
                # [x_min, y_min, x_max, y_max]
                x_min, y_min, x_max, y_max = rec_boxes[i]
                xs = [x_min, x_max]
                ys = [y_min, y_max]
            elif rec_polys is not None and i < len(rec_polys):
                # [[x1,y1],[x2,y2],[x3,y3],[x4,y4]]
                poly = rec_polys[i]
                xs = [p[0] for p in poly]
                ys = [p[1] for p in poly]
            elif dt_polys is not None and i < len(dt_polys):
                poly = dt_polys[i]
                xs = [p[0] for p in poly]
                ys = [p[1] for p in poly]
            else:
                # No usable box; skip
                continue

            x_min = min(xs)
            y_center = sum(ys) / len(ys)

            all_boxes.append((y_center, x_min, text))

    return all_boxes


def extract_text_from_image(image_path, file_id="output"):
    """
    Run OCR on the given image using PaddleOCR 3.x pipeline and
    save the recognized text as a Word file.

    Parameters
    ----------
    image_path : str or Path
        Path to the input image (JPG / PNG / etc.).
    file_id : str
        Logical ID used to name the output .docx file.
        The final file will be:
            E:\\ImageTextSqubix\\temp\\<file_id>.docx

    Returns
    -------
    str
        Full path to the generated .docx file (as a string).
    """
    image_path = Path(image_path)

    if not image_path.is_file():
        raise FileNotFoundError(f"Image not found: {image_path}")

    # Ensure output directory exists
    TEMP_TEXT_DIR.mkdir(parents=True, exist_ok=True)

    # ------------------------------------------------------------------
    # 1) Run PaddleOCR 3.x pipeline
    #    - predict() is the recommended API in v3
    # ------------------------------------------------------------------
    result_list = ocr.predict(str(image_path))

    # ------------------------------------------------------------------
    # 2) Convert raw result into our own list of boxes
    # ------------------------------------------------------------------
    all_boxes = _collect_boxes_from_result(result_list)

    # ------------------------------------------------------------------
    # 3) Group into visual lines (top-to-bottom, left-to-right)
    # ------------------------------------------------------------------
    lines = []
    if all_boxes:
        # Sort top-to-bottom, then left-to-right
        all_boxes.sort(key=lambda t: (t[0], t[1]))

        current_line = []
        current_y = None
        Y_THRESHOLD = 15  # pixels tolerance for "same line"

        for y_center, x_min, text in all_boxes:
            if current_y is None:
                # First word
                current_y = y_center
                current_line.append((x_min, text))
            elif abs(y_center - current_y) <= Y_THRESHOLD:
                # Same row as previous
                current_line.append((x_min, text))
            else:
                # New visual row -> flush old one
                current_line.sort(key=lambda t: t[0])
                lines.append(" ".join(t for _, t in current_line))

                current_line = [(x_min, text)]
                current_y = y_center

        # Flush last line
        if current_line:
            current_line.sort(key=lambda t: t[0])
            lines.append(" ".join(t for _, t in current_line))

    # ------------------------------------------------------------------
    # 4) (Optional) post-cleaning – you can filter watermarks here
    # ------------------------------------------------------------------
    # Example:
    # lines = [ln for ln in lines if "shutterstock" not in ln.lower()]

    # ------------------------------------------------------------------
    # 5) Write DOCX
    # ------------------------------------------------------------------
    docx_path = TEMP_TEXT_DIR / f"{file_id}.docx"
    doc = Document()

    if lines:
        for line in lines:
            doc.add_paragraph(line)
    else:
        # To help debug, you can also write a placeholder
        doc.add_paragraph("[No text detected by OCR]")

    doc.save(docx_path)

    return str(docx_path)
